import streamlit as st
import os
import torch
from transformers import GPT2LMHeadModel, GPT2Tokenizer
from PyPDF2 import PdfReader
from docx import Document
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import base64
from docx2pdf import convert
import pythoncom
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random


@st.cache_resource
def load_model():
    model_name = "gpt2"
    tokenizer = GPT2Tokenizer.from_pretrained(model_name)
    
    if tokenizer.pad_token is None:
        tokenizer.pad_token = tokenizer.eos_token
        
    model = GPT2LMHeadModel.from_pretrained(model_name)
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model = model.to(device)
    tokenizer.padding_side = 'left'
    model.config.pad_token_id = tokenizer.eos_token_id
    return tokenizer, model

def extract_information_with_gpt(text, tokenizer, model, option):
    try:
        if option == "Surat Tugas Dosen":
            prompt = f"""
            Anda adalah asisten AI yang bertugas untuk mengekstrak informasi penting dari dokumen pengajuan.
            Berikut adalah teks dokumen: {text}
            Format hasil ekstraksi:
            Nama Lengkap: [nama]
            NIM: [nim]
            Perlombaan: [perlombaan]
            Tanggal: [tanggal]
            Lokasi: [lokasi]
            Dosen Pembimbing (nama lengkap dengan gelarnya): [dosen pembimbing]
            NIP Dospem: [nip]
            Jurusan: [jurusan]
            """
        else:  # Pengajuan Dispensasi
            prompt = f"""
            Anda adalah asisten AI yang bertugas untuk mengekstrak informasi penting dari dokumen dispensasi.
            Berikut adalah teks dokumen: {text}
            Format hasil ekstraksi:
            Nama Lengkap: [nama]
            NIM (yang bernilai angka): [nim]
            Jurusan: [jurusan]
            Alasan Dispensasi: [alasan]
            Tanggal Mulai: [tanggal mulai]
            Tanggal Selesai: [tanggal selesai]
            """
        
        inputs = tokenizer.encode(
            prompt,
            return_tensors='pt',
            max_length=825,
            truncation=True
        )
        
        outputs = model.generate(
            inputs,
            max_new_tokens=200,
            do_sample=True,
            temperature=0.7,
            top_p=0.9,
            num_return_sequences=1,
            eos_token_id=tokenizer.eos_token_id
        )
        
        response = tokenizer.decode(outputs[0], skip_special_tokens=True)
        
        extracted_info = {}
        for line in response.split('\n'):
            if ':' in line:
                key, value = line.split(':', 1)
                if value.strip():  # Hanya ambil nilai yang tidak kosong
                    extracted_info[key.strip()] = value.strip()
        
        return extracted_info
        
    except Exception as e:
        print(f"Error dalam ekstraksi: {str(e)}")
        return {}

def buat_template_surat(nama_dosen, nip_dosen, nama_mahasiswa, nim_mahasiswa, jurusan, nama_lomba, tanggal_lomba, tanggal_surat, dekan, nip_dekan):
    doc = Document()
    
    title = doc.add_heading('SURAT TUGAS', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nomor_surat = doc.add_paragraph('Nomor: 06274/UN10.F1501/B/KM/2024')
    nomor_surat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')
    
    content = doc.add_paragraph()
    content.add_run('Bersama ini Dekan Fakultas Ilmu Komputer Universitas Brawijaya menugaskan kepada personalia:').bold = True
    
    doc.add_paragraph(f'Nama: {nama_dosen}')
    doc.add_paragraph(f'NIP/NIK/NIDN: {nip_dosen}')
    
    doc.add_paragraph(f'Untuk menjadi dosen pembimbing {nama_lomba} yang diselenggarakan tanggal {tanggal_lomba}.')
    
    mahasiswa_section = doc.add_paragraph('Mahasiswa yang dibimbing adalah sebagai berikut:')
    mahasiswa_section.add_run('\n')
    doc.add_paragraph(f'1. {nama_mahasiswa} ({nim_mahasiswa}) - {jurusan}')
    
    closing = doc.add_paragraph()
    closing.add_run('Demikian Surat Tugas ini dibuat untuk dilaksanakan dengan sebaik-baiknya dan penuh rasa tanggung jawab.').italic = True
    
    doc.add_paragraph('\n')
    signature = doc.add_paragraph()
    signature.add_run(f'{tanggal_surat}\nDekan,\n\n\n{dekan}\nNIP {nip_dekan}')
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    
    folder_name = "surat_tugas"
    file_name = f'{folder_name}/surat_tugas_{nama_mahasiswa.lower().replace(" ", "_")}.docx'
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)
    doc.save(file_name)
    return file_name

def buat_template_surat_dispensasi(nama_mahasiswa, nim_mahasiswa, jurusan, alasan_dispensasi, tanggal_mulai, tanggal_selesai, tanggal_surat, dekan, nip_dekan):
    doc = Document()
    
    # Judul surat
    title = doc.add_heading('SURAT DISPENSASI', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Nomor surat
    nomor_surat = doc.add_paragraph('Nomor: ' + generate_nomor_surat())
    nomor_surat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    
    # Isi surat
    content = doc.add_paragraph()
    content.add_run('Yang bertanda tangan di bawah ini, Dekan Fakultas Ilmu Komputer Universitas Brawijaya, memberikan dispensasi kepada mahasiswa:').bold = True
    
    doc.add_paragraph(f'Nama: {nama_mahasiswa}')
    doc.add_paragraph(f'NIM: {nim_mahasiswa}')
    doc.add_paragraph(f'Jurusan: {jurusan}')
    
    doc.add_paragraph(f'Untuk mendapatkan dispensasi {alasan_dispensasi} yang berlaku mulai tanggal {tanggal_mulai} sampai dengan {tanggal_selesai}.')
    
    doc.add_paragraph('Demikian surat dispensasi ini dibuat untuk dipergunakan sebagaimana mestinya.')
    
    doc.add_paragraph('\n')
    
    # Tanda tangan
    signature = doc.add_paragraph()
    signature.add_run(f'{tanggal_surat}\nDekan,\n\n\n{dekan}\nNIP {nip_dekan}')
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    folder_name = "surat_dispen"    
    file_name = f'{folder_name}/surat_dispensasi_{nama_mahasiswa.lower().replace(" ", "_")}.docx'
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)
    doc.save(file_name)
    return file_name

def generate_nomor_surat():
    # Fungsi untuk menghasilkan nomor surat yang unik
    tahun = datetime.now().year
    nomor = f"{random.randint(1000, 9999)}/UN10.F1501/AK/KM/{tahun}"
    return nomor


def convert_docx_to_pdf(docx_filename, pdf_filename):
    convert(docx_filename, pdf_filename)


def read_pdf(file_path):
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""

def save_uploaded_file(uploaded_file, option):
    folder_name = "STD" if option == "Surat Tugas Dosen" else "Dispen"
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)
        
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_name, file_extension = os.path.splitext(uploaded_file.name)
    new_filename = f"{file_name}_{timestamp}{file_extension}"
    file_path = os.path.join(folder_name, new_filename)
    
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getvalue())
    return file_path

def display_pdf(pdf_file):
    with open(pdf_file, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="700" height="1000" type="application/pdf"></iframe>'
    return pdf_display

def main():
    pythoncom.CoInitialize()
    st.set_page_config(page_title="Aplikasi Pengajuan Dokumen Kemahasiswaan", layout="wide")
    st.title("Aplikasi Pengajuan Dokumen Kemahasiswaan")
    
    st.sidebar.title("Menu")
    option = st.sidebar.radio("Pilih Jenis Pengajuan", ("Surat Tugas Dosen", "Pengajuan Dispensasi"))
    
    st.header(f"Form Pengajuan {option}")
    uploaded_file = st.file_uploader("Upload dokumen (PDF):", type=['pdf'])
    
    if uploaded_file is not None:
        file_details = {
            "Nama File": uploaded_file.name,
            "Tipe File": uploaded_file.type,
            "Ukuran": f"{uploaded_file.size / 1024:.2f} KB"
        }
        # st.write("Detail File:")
        # st.json(file_details)
        
        if st.button("Proses Dokumen"):
            with st.spinner('Memproses dokumen...'):
                try:
                    saved_file_path = save_uploaded_file(uploaded_file, option)
                    st.success(f"File berhasil disimpan di folder {os.path.dirname(saved_file_path)}")
                    
                    if uploaded_file.type == "application/pdf":
                        text_content = read_pdf(saved_file_path)
                        if text_content:
                            tokenizer, model = load_model()
                            extracted_info = extract_information_with_gpt(text_content, tokenizer, model, option)
                            
                            # st.write("Raw extracted info:", extracted_info)
                            
                            if option == "Surat Tugas Dosen":
                                docx_file = buat_template_surat(
                                    nama_dosen=extracted_info.get('NAMA DOSEN PEMBIMBING', ''),
                                    nip_dosen='0012056904',
                                    nama_mahasiswa=extracted_info.get('NAMA  LENGKAP', ''),
                                    nim_mahasiswa=extracted_info.get('NIM', ''),
                                    jurusan=extracted_info.get('FAKULTAS', ''),
                                    nama_lomba=extracted_info.get('NAMA PERLOMBAAN', ''),
                                    tanggal_lomba=extracted_info.get('TANGGAL PELAKSANAAN', ''),
                                    tanggal_surat=datetime.now().strftime("%d %B %Y"),
                                    dekan='Wayan Firdaus Mahmudy',
                                    nip_dekan='196012301986011001'
                                )
                            else:  # Pengajuan Dispensasi
                                print(extracted_info)
                                docx_file = buat_template_surat_dispensasi(
                                    nama_mahasiswa=extracted_info.get('NAMA  LENGKAP', ''),
                                    nim_mahasiswa=extracted_info.get('NIM', ''),
                                    jurusan=extracted_info.get('FAKULTAS', ''),
                                    alasan_dispensasi=extracted_info.get('NAMA KEGIATAN', ''),
                                    tanggal_mulai=extracted_info.get('TANGGAL MULAI', ''),
                                    tanggal_selesai=extracted_info.get('TANGGAL SELESAI', ''),
                                    tanggal_surat=datetime.now().strftime("%d %B %Y"),
                                    dekan='Wayan Firdaus Mahmudy',
                                    nip_dekan='196012301986011001'
                                )
                            
                            pdf_file = docx_file.replace('.docx', '.pdf')
                            convert_docx_to_pdf(docx_file, pdf_file)
                            
                            tab1, tab2 = st.tabs(["Hasil Ekstraksi", "Preview Dokumen"])
                            with tab1:
                                st.subheader("Informasi yang Diekstrak:")
                                for key in extracted_info.keys():
                                    st.write(f"**{key}:** {extracted_info[key]}")
                                    
                            with tab2:
                                st.subheader("Dokumen yang Dihasilkan:")
                                html_code = display_pdf(pdf_file)
                                st.markdown(html_code, unsafe_allow_html=True)
                        else:
                            st.error("Tidak dapat membaca teks dari PDF")
                    else:
                        st.error("Format file tidak didukung")
                except Exception as e:
                    st.error(f"Terjadi kesalahan: {str(e)}")

if __name__ == "__main__":
    main()
